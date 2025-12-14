"""DOCX file processor."""

from docx import Document


def process_docx(file_path: str) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Extracted text content
    """
    doc = Document(file_path)
    text = ""
    
    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text += "\n"
    
    return text
